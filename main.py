from flask import Flask, request, flash, redirect, url_for, render_template, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_wtf import FlaskForm
from wtforms import FileField, SubmitField
from wtforms.validators import DataRequired
from werkzeug.utils import secure_filename
import os
import logging
from docx import Document
import win32com.client
from openpyxl import Workbook, load_workbook

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
app.config['SECRET_KEY'] = "random string"
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['WTF_CSRF_ENABLED'] = True

logging.basicConfig(level=logging.DEBUG)

db = SQLAlchemy(app)

class Students(db.Model):
   id = db.Column('student_id', db.Integer, primary_key=True)
   name = db.Column(db.String(100))
   city = db.Column(db.String(50))
   addr = db.Column(db.String(200))
   pin = db.Column(db.String(10))

   def __init__(self, name, city, addr, pin):
      self.name = name
      self.city = city
      self.addr = addr
      self.pin = pin

class UploadForm(FlaskForm):
    file = FileField('File', validators=[DataRequired()])
    submit = SubmitField('Upload')

@app.route('/', methods=['POST', 'GET'])
def show_all():
    students_list = Students.query.all()
    search_value = request.form.get('searchValue')
    delete_value = request.form.get('deleteValue')

    if request.method == 'POST':
        if search_value:
            students_list = Students.query.filter(Students.name.like(f"%{search_value}%")).all()
        else:
            students_list = Students.query.all()
    else:
        students_list = Students.query.all()

    return render_template('pages/show_all.html', students=students_list)

@app.route('/new', methods=['GET', 'POST'])
def new():
   if request.method == 'POST':
      if not request.form['name'] or not request.form['city'] or not request.form['addr']:
         flash('Please enter all the fields', 'error')
      else:
         student = Students(request.form['name'], request.form['city'],
                            request.form['addr'], request.form['pin'])
         
         db.session.add(student)
         db.session.commit()
         flash('Record was successfully added')
         return redirect(url_for('show_all'))
   return render_template('pages/new.html')

@app.route('/update/<int:student_id>', methods=['POST'])
def update_student(student_id):
    student = Students.query.get(student_id)
    if student:
        student.name = request.form['name']
        student.city = request.form['city']
        student.addr = request.form['addr']
        student.pin = request.form['pin']
        db.session.commit()
        flash('Student successfully updated')
    else:
        flash('Student not found')
    return redirect(url_for('show_all'))

@app.route('/delete/<int:student_id>', methods=['POST'])
def delete_student(student_id):
    student = Students.query.get(student_id)
    if student:
        db.session.delete(student)
        db.session.commit()
        flash('Student successfully deleted')
    else:
        flash('Student not found')
    return redirect(url_for('show_all'))

@app.route('/upload', methods=['GET', 'POST'])
def upload():
    form = UploadForm() # tạo 1 biểu mẫu upload
    if form.validate_on_submit(): # kiểm tra xem form đã được submit và hợp lệ chưa
        file = form.file.data # lấy dữ liệu file được upload
        filename = secure_filename(file.filename) # lấy file name, dùng secure_filename để đảm bảo an toàn cho file tránh các đoạn mã nguy hiểm ký tự đặc biệt
        
        if not os.path.exists(app.config['UPLOAD_FOLDER']): # kiểm tra xem thư mục lưu đã được tạo chưa, chưa sẽ tạo mới
            os.makedirs(app.config['UPLOAD_FOLDER'])

        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename) # đường dẫn đầy đủ cho file được upload
        file.save(file_path) # lưu file upload vào đường dẫn đã được dựng

        if filename.endswith('.csv'): 
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            except UnicodeDecodeError as e:
                flash(f'Error decoding CSV file: {e}')
                return redirect(url_for('upload'))

            lines = content.split('\n') # dùng split để tách từng dòng
            for line in lines:
                parts = line.strip().split(',') # strip để xóa khoảng trắng và tách từng part cách nhau 1 dấu ","
                if len(parts) == 4: # kiểm tra độ dài list có bằng 4 không
                    name, city, addr, pin = parts # gán các giá trị của list cho các biến
                    student = Students(name.strip(), city.strip(), addr.strip(), pin.strip())
                    db.session.add(student)
            db.session.commit()
            flash('CSV file successfully uploaded and processed')

        elif filename.endswith('.docx'):
            try:
                doc = Document(file_path) # chuyển đổi file docx từ file_path thành đối tượng Document
                for paragraph in doc.paragraphs: # lặp qua các đoạn văn bản có trong file sau khi chuyển đổi qua Document 
                    parts = paragraph.text.strip().split(',') # paragraph.text trả về dạng text của đoạn văn bản hiện tại
                    if len(parts) == 4:
                        name, city, addr, pin = parts
                        student = Students(name.strip(), city.strip(), addr.strip(), pin.strip())
                        db.session.add(student)
                db.session.commit()
                flash('Word document successfully uploaded and processed')

            except Exception as e:
                flash(f'Error processing Word document: {str(e)}')

        elif filename.endswith('.doc'):
            try:
                import pythoncom
                pythoncom.CoInitialize()
                word_app = win32com.client.Dispatch("Word.Application") # khởi tạo 1 phiên làm việc với Word
                word_app.Visible = True

                doc = word_app.Documents.Open(file_path)
                content = doc.Content.Text
                lines = content.split('\n')
                for line in lines:
                    parts = line.strip().split(',')
                    if len(parts) == 4:
                        name, city, addr, pin = parts
                        student = Students(name.strip(), city.strip(), addr.strip(), pin.strip())
                        db.session.add(student)
                db.session.commit()
                flash('Word document successfully uploaded and processed')

                doc.Close()
                word_app.Quit()
                
            except Exception as e:
                flash(f'Error processing Word document: {str(e)}')
        
        elif file_path.endswith('.xlsx'):
            try:
                wb = load_workbook(file) # Sử dụng load_workbook từ openpyxl để tải workbook từ file Excel
                ws = wb.active

                # Sử dụng iter_rows để lặp qua từng hàng trong worksheet và nhập dữ liệu vào database.
                # values_only=True: nó sẽ trả về giá trị của các ô thay vì các đối tượng ô 
                for row in ws.iter_rows(min_row=2, values_only=True): # bắt đầu từ dòng thứ 2 do dòng 1 là tiêu đề
                    name, city, addr, pin = row
                    student = Students(name=name, city=city, addr=addr, pin=pin)
                    db.session.add(student)

                db.session.commit()    
                flash('Excel document uploaded successfully')
            except Exception as e:
                flash(f'Error processing Excel document: {str(e)}')

        else:
            flash('Unsupported file format. Please upload a CSV or Word or Excel document.')
        
        return redirect(url_for('show_all'))

    return render_template('pages/upload.html', form=form)

@app.route('/export_docx', methods=['GET'])
def export_docx():
    students = Students.query.all()

    doc = Document() # tạo 1 Document mới
    doc.add_heading('Student List') # thêm tiêu đề cho docs

    # thêm dữ liệu vào docs
    for student in students:
        doc.add_paragraph(f'{student.name}, {student.city}, {student.addr}, {student.pin}')
    
    # lưu tài liệu xuống file
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'student_list.docx')
    doc.save(file_path)

    # trả về file cho người dùng tải xuống
    return send_file(file_path, as_attachment=True,download_name='student_list.docx')

@app.route('/export_excel', methods=['GET'])
def export_excel():
    students = Students.query.all()

    # Tạo một Workbook mới
    wb = Workbook()
    ws = wb.active
    ws.title = "Student List"

    # Tiêu đề cột
    columns = ["Name", "City", "Address", "PIN"]
    ws.append(columns)

    # Thêm dữ liệu từ database vào worksheet
    for student in students:
        ws.append([student.name, student.city, student.addr, student.pin])

    # Điều chỉnh độ rộng cột dựa trên giá trị dài nhất
    for column_cells in ws.columns:
        length = max(len(str(cell.value)) for cell in column_cells)
        column_letter = column_cells[0].column_letter
        ws.column_dimensions[column_letter].width = length + 2  # Thêm một chút đệm để cột không quá khít

    # Lưu workbook xuống file
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'student_list.xlsx')
    wb.save(file_path)

    # Trả về file cho người dùng tải xuống
    return send_file(file_path, as_attachment=True, download_name='student_list.xlsx')

if __name__ == '__main__':
   with app.app_context():
      db.create_all()
   app.run(debug=True)
